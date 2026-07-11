#!/usr/bin/env python
"""Inspect a Word template and write a compact JSON report.

Usage:
  python inspect_docx_template.py template.docx --out template_report.json
"""
from __future__ import annotations
import sys

import argparse
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


def pt(value):
    return None if value is None else round(value.pt, 2)


def style_summary(style):
    """Summarize a paragraph or character style.

    Table / numbering styles do not expose .font / .paragraph_format and would
    raise AttributeError if probed the same way as paragraph styles, so we
    guard each access defensively.
    """
    info = {
        "name": style.name,
        "style_id": style.style_id,
        "type": str(style.type),
    }
    font = getattr(style, "font", None)
    if font is not None:
        try:
            info["font_name"] = font.name
            info["font_size_pt"] = pt(font.size)
            info["bold"] = font.bold
        except Exception:
            pass
    pf = getattr(style, "paragraph_format", None)
    if pf is not None:
        try:
            info["line_spacing"] = str(pf.line_spacing)
            info["space_before_pt"] = pt(pf.space_before)
            info["space_after_pt"] = pt(pf.space_after)
            info["first_line_indent_pt"] = pt(pf.first_line_indent)
            info["left_indent_pt"] = pt(pf.left_indent)
        except Exception:
            pass
    try:
        xml = style._element.xml
        info["has_numbering"] = "numPr" in xml
        info["outline_level"] = "outlineLvl" in xml
    except Exception:
        pass
    return info


def paragraph_summary(paragraph, idx):
    runs = []
    for run in paragraph.runs[:8]:
        if not run.text:
            continue
        color = None
        if run.font.color and run.font.color.rgb:
            color = str(run.font.color.rgb)
        runs.append(
            {
                "text": run.text[:80],
                "font": run.font.name,
                "size_pt": pt(run.font.size),
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "color": color,
                "superscript": run.font.superscript,
            }
        )
    return {
        "index": idx,
        "style": paragraph.style.name,
        "text": paragraph.text[:200],
        "alignment": str(paragraph.alignment),
        "has_drawing": "<w:drawing" in paragraph._element.xml,
        "has_omath": "m:oMath" in paragraph._element.xml,
        "runs": runs,
    }


def table_summary(table, idx):
    preview = []
    for row in table.rows[:3]:
        preview.append([cell.text[:80] for cell in row.cells[:6]])
    return {
        "index": idx,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "preview": preview,
    }


def is_heading_style(name: str) -> bool:
    """Match both English ('Heading 1') and Chinese ('标题 1') heading styles."""
    return name.startswith("Heading") or "标题" in name


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("template")
    parser.add_argument("--out", default=None)
    args = parser.parse_args()

    path = Path(args.template)
    doc = Document(path)

    used_style_names = {p.style.name for p in doc.paragraphs}

    # 关键修复：模板中预定义但 body 内无段落使用的样式 —— 例如 "论文正文"、
    # "参考文献条目" —— 才是 AI 真正需要识别并映射的目标。原实现只导出
    # name in used 的样式，会把这些关键样式整体丢掉。
    relevant_types = {WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER}
    styles_export = []
    for style in doc.styles:
        try:
            stype = style.type
        except Exception:
            continue
        if stype not in relevant_types:
            continue
        try:
            summary = style_summary(style)
        except Exception:
            continue
        summary["used_in_body"] = style.name in used_style_names
        styles_export.append(summary)

    report = {
        "file": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "styles_used_in_body": sorted(used_style_names),
        "styles": styles_export,
        "paragraphs_first_120": [
            paragraph_summary(p, i) for i, p in enumerate(doc.paragraphs[:120])
        ],
        "heading_paragraphs": [
            paragraph_summary(p, i)
            for i, p in enumerate(doc.paragraphs)
            if is_heading_style(p.style.name)
        ],
        "tables": [table_summary(t, i) for i, t in enumerate(doc.tables)],
    }

    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        if "word/document.xml" in names:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
            report["document_xml_hints"] = {
                "hyperlink_count": xml.count("<w:hyperlink"),
                "bookmark_count": xml.count("<w:bookmarkStart"),
                "omath_count": xml.count("<m:oMath"),
                "blue_hyperlink_color": (
                    'w:color w:val="0563C1"' in xml
                    or 'w:color w:val="0000FF"' in xml
                ),
            }
        if "word/numbering.xml" in names:
            numbering = zf.read("word/numbering.xml").decode("utf-8", errors="ignore")
            num_ids = sorted(
                {int(m.group(1)) for m in re.finditer(r'w:numId w:val="(\d+)"', numbering)}
            )
            abstract_ids = sorted(
                {int(m.group(1)) for m in re.finditer(r'w:abstractNumId w:val="(\d+)"', numbering)}
            )
            report["numbering"] = {
                "xml_length": len(numbering),
                "num_ids": num_ids,
                "abstract_num_ids": abstract_ids,
            }

    out = Path(args.out) if args.out else path.with_suffix(".template-report.json")
    out.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(out)
    return 0


if __name__ == "__main__":
    for _s in (sys.stdout, sys.stderr):
        if hasattr(_s, "reconfigure"):
            _s.reconfigure(encoding="utf-8", errors="replace")

    raise SystemExit(main())
