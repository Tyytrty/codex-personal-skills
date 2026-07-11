#!/usr/bin/env python
"""Validate a reconstructed DOCX/PDF against common template-conversion failures.

Usage:
  python validate_docx_conversion.py final.docx --pdf final.pdf --out validation.json
  python validate_docx_conversion.py final.docx --ordered-term 绪论 --ordered-term 参考文献
  python validate_docx_conversion.py final.docx --template template.docx --protected-until 中 文 摘 要
"""
from __future__ import annotations
import sys

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DEFAULT_PLACEHOLDERS = [
    "李四",
    "王五",
    "张三",
    "lorem ipsum",
    "Lorem ipsum",
    "为了提高本科生学位论文的质量",
    "此处为论文题目的英文翻译",
    "题目（一般不宜超过25字",
    "右键更新域或由自动最终化脚本更新目录",
]


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def paragraph_style_name(paragraph) -> str:
    try:
        return paragraph.style.name
    except Exception:
        return ""


def paragraph_style_id(paragraph) -> str:
    try:
        return paragraph.style.style_id
    except Exception:
        return ""


def xml_bool(run_rpr, tag_name: str):
    if run_rpr is None:
        return None
    node = run_rpr.find(qn(f"w:{tag_name}"))
    if node is None:
        return None
    val = node.get(qn("w:val"))
    return False if val in {"0", "false", "False"} else True


def xml_child_val(parent, tag_name: str) -> str | None:
    if parent is None:
        return None
    node = parent.find(qn(f"w:{tag_name}"))
    if node is None:
        return None
    return node.get(qn("w:val"))


def run_format_signature(run) -> dict:
    rpr = run._element.rPr
    fonts = {}
    if rpr is not None:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            fonts = {key.split("}")[-1]: value for key, value in rfonts.attrib.items()}
    return {
        "fonts": fonts,
        "size": xml_child_val(rpr, "sz"),
        "bold": xml_bool(rpr, "b"),
        "italic": xml_bool(rpr, "i"),
        "underline": xml_child_val(rpr, "u"),
        "color": xml_child_val(rpr, "color"),
        "vert_align": xml_child_val(rpr, "vertAlign"),
    }


def paragraph_format_signature(paragraph) -> dict:
    return {
        "style": paragraph_style_name(paragraph),
        "style_id": paragraph_style_id(paragraph),
        "alignment": str(paragraph.alignment),
        "run_count": len(paragraph.runs),
        "run_formats": [run_format_signature(run) for run in paragraph.runs],
    }


def paragraphs_before_marker(doc: Document, marker: str) -> tuple[list, bool]:
    needle = normalize(marker)
    paragraphs = []
    for paragraph in doc.paragraphs:
        if needle and needle in normalize(paragraph.text):
            return paragraphs, True
        paragraphs.append(paragraph)
    return paragraphs, False


def compare_protected_front_matter(
    template_path: Path,
    final_doc: Document,
    *,
    protected_until: str,
    max_diffs: int,
) -> tuple[dict, list[str]]:
    template_doc = Document(str(template_path))
    template_paragraphs, template_found = paragraphs_before_marker(
        template_doc, protected_until
    )
    final_paragraphs, final_found = paragraphs_before_marker(final_doc, protected_until)

    diffs = []
    failures: list[str] = []
    if not template_found:
        failures.append(f"protected marker missing in template: {protected_until}")
    if not final_found:
        failures.append(f"protected marker missing in final docx: {protected_until}")
    if len(template_paragraphs) != len(final_paragraphs):
        failures.append(
            "protected front matter paragraph count changed: "
            f"{len(template_paragraphs)} != {len(final_paragraphs)}"
        )

    for idx, (template_p, final_p) in enumerate(
        zip(template_paragraphs, final_paragraphs)
    ):
        template_sig = paragraph_format_signature(template_p)
        final_sig = paragraph_format_signature(final_p)
        if template_sig == final_sig:
            continue
        if len(diffs) < max_diffs:
            diffs.append(
                {
                    "index": idx,
                    "template_text": template_p.text[:120],
                    "final_text": final_p.text[:120],
                    "template_format": template_sig,
                    "final_format": final_sig,
                }
            )

    if diffs:
        failures.append("protected front matter formatting changed")

    report = {
        "template": str(template_path),
        "protected_until": protected_until,
        "template_marker_found": template_found,
        "final_marker_found": final_found,
        "template_paragraph_count": len(template_paragraphs),
        "final_paragraph_count": len(final_paragraphs),
        "diff_count_sampled": len(diffs),
        "diffs": diffs,
    }
    return report, failures


def collect_docx_report(
    docx_path: Path,
    *,
    template_path: Path | None,
    protected_until: str | None,
    protected_max_diffs: int,
    placeholders: list[str],
    ordered_terms: list[str],
    required_headings: list[str],
    forbidden_header_terms: list[str],
    min_images: int,
    min_tables: int,
) -> tuple[dict, list[str]]:
    doc = Document(str(docx_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    compact_text = normalize(text)
    failures: list[str] = []

    placeholder_hits = [term for term in placeholders if term and term in text]
    if placeholder_hits:
        failures.append("template placeholder text remains")

    order_positions = {term: compact_text.find(normalize(term)) for term in ordered_terms}
    if ordered_terms:
        previous = -1
        for term in ordered_terms:
            pos = order_positions[term]
            if pos < 0:
                failures.append(f"ordered term missing: {term}")
            elif pos <= previous:
                failures.append(f"ordered term out of order: {term}")
            previous = pos

    heading_styles = {"Heading 1", "Heading 2", "Heading 3", "标题 1", "标题 2", "标题 3"}
    headings = [
        {"text": p.text.strip(), "style": paragraph_style_name(p)}
        for p in doc.paragraphs
        if paragraph_style_name(p) in heading_styles
    ]
    heading_compact = {normalize(item["text"]): item for item in headings}
    missing_heading_terms = [
        term for term in required_headings if normalize(term) not in heading_compact
    ]
    if missing_heading_terms:
        failures.append("required headings are not heading-styled")

    section_headers = []
    forbidden_header_hits = []
    for idx, section in enumerate(doc.sections):
        header_text = "\n".join(p.text for p in section.header.paragraphs).strip()
        section_headers.append({"section": idx, "text": header_text})
        for term in forbidden_header_terms:
            if term and term in header_text:
                forbidden_header_hits.append({"section": idx, "term": term, "text": header_text})
        if forbidden_header_terms:
            # Also catch stale header references that may render through Word even
            # when python-docx sees little visible text.
            refs = section._sectPr.findall(qn("w:headerReference"))
            if refs and not header_text:
                section_headers[-1]["header_reference_count"] = len(refs)
    if forbidden_header_hits:
        failures.append("forbidden back-matter header inherited by body section")

    image_count = len(doc.inline_shapes)
    table_count = len(doc.tables)
    if image_count < min_images:
        failures.append(f"image count below minimum: {image_count} < {min_images}")
    if table_count < min_tables:
        failures.append(f"table count below minimum: {table_count} < {min_tables}")

    report = {
        "docx": str(docx_path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": table_count,
        "image_count": image_count,
        "placeholder_hits": placeholder_hits,
        "ordered_terms": order_positions,
        "heading_count": len(headings),
        "missing_heading_terms": missing_heading_terms,
        "section_headers": section_headers,
        "forbidden_header_hits": forbidden_header_hits,
    }
    if template_path and protected_until:
        protected_report, protected_failures = compare_protected_front_matter(
            template_path,
            doc,
            protected_until=protected_until,
            max_diffs=protected_max_diffs,
        )
        report["protected_front_matter"] = protected_report
        failures.extend(protected_failures)
    return report, failures


def collect_pdf_report(pdf_path: Path, placeholders: list[str]) -> tuple[dict, list[str]]:
    failures: list[str] = []
    try:
        import fitz
    except ImportError:
        return {"pdf": str(pdf_path), "error": "PyMuPDF not installed"}, ["pdf validation unavailable"]

    with fitz.open(pdf_path) as pdf:
        page_texts = [page.get_text() for page in pdf]
    text = "\n".join(page_texts)
    placeholder_hits = [term for term in placeholders if term and term in text]
    if placeholder_hits:
        failures.append("template placeholder text remains in PDF")
    report = {
        "pdf": str(pdf_path),
        "page_count": len(page_texts),
        "placeholder_hits": placeholder_hits,
    }
    return report, failures


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx")
    parser.add_argument("--pdf", default=None)
    parser.add_argument("--template", default=None)
    parser.add_argument(
        "--protected-until",
        default=None,
        help="Marker that starts generated content; compare template formatting before it",
    )
    parser.add_argument("--protected-max-diffs", type=int, default=20)
    parser.add_argument("--out", default=None)
    parser.add_argument("--placeholder", action="append", default=[])
    parser.add_argument("--no-default-placeholders", action="store_true")
    parser.add_argument("--ordered-term", action="append", default=[])
    parser.add_argument("--required-heading", action="append", default=[])
    parser.add_argument("--forbidden-header", action="append", default=[])
    parser.add_argument("--min-images", type=int, default=0)
    parser.add_argument("--min-tables", type=int, default=0)
    args = parser.parse_args()
    if bool(args.template) != bool(args.protected_until):
        parser.error("--template and --protected-until must be used together")

    docx_path = Path(args.docx)
    placeholders = list(args.placeholder)
    if not args.no_default_placeholders:
        placeholders = DEFAULT_PLACEHOLDERS + placeholders

    docx_report, failures = collect_docx_report(
        docx_path,
        template_path=Path(args.template) if args.template else None,
        protected_until=args.protected_until,
        protected_max_diffs=args.protected_max_diffs,
        placeholders=placeholders,
        ordered_terms=args.ordered_term,
        required_headings=args.required_heading,
        forbidden_header_terms=args.forbidden_header,
        min_images=args.min_images,
        min_tables=args.min_tables,
    )
    report = {"docx": docx_report}

    if args.pdf:
        pdf_report, pdf_failures = collect_pdf_report(Path(args.pdf), placeholders)
        report["pdf"] = pdf_report
        failures.extend(pdf_failures)

    status = "PASS" if not failures else "FAIL"
    report["status"] = status
    report["failures"] = failures

    output = json.dumps(report, ensure_ascii=False, indent=2)
    if args.out:
        Path(args.out).write_text(output, encoding="utf-8")
    print(f"STATUS: {status}")
    if failures:
        for failure in failures:
            print(f"- {failure}")
    if args.out:
        print(args.out)
    return 0 if status == "PASS" else 1


if __name__ == "__main__":
    for _s in (sys.stdout, sys.stderr):
        if hasattr(_s, "reconfigure"):
            _s.reconfigure(encoding="utf-8", errors="replace")

    raise SystemExit(main())
